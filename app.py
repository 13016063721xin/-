import base64
import io

from fastapi import FastAPI, File, UploadFile, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def decode_data_url(data_url: str) -> bytes:
    if not data_url or "," not in data_url:
        raise ValueError("Invalid data URL")
    _, encoded = data_url.split(",", 1)
    return base64.b64decode(encoded)


@app.post("/api/img2doc")
async def convert_image_to_docx(file: UploadFile = File(...)):
    image_data = await file.read()
    image_stream = io.BytesIO(image_data)

    doc = Document()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.add_picture(image_stream, width=Inches(6.0))

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    original_name = file.filename.rsplit(".", 1)[0] if file.filename and "." in file.filename else "image"
    headers = {
        "Content-Disposition": f'attachment; filename="{original_name}_converted.docx"'
    }

    return StreamingResponse(
        doc_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


@app.post("/upload")
async def export_mixed_word(payload=Body(...)):
    try:
      items = payload.get("items", [])
      if not isinstance(items, list) or not items:
          return JSONResponse(
              status_code=400,
              content={"error": "items must be a non-empty array"}
          )

      doc = Document()

      for item in items:
          item_type = item.get("type")

          if item_type == "text":
              content = str(item.get("content", "")).strip()
              if content:
                  p = doc.add_paragraph()
                  p.add_run(content)

          elif item_type == "image":
              data_url = item.get("content", "")
              if not data_url:
                  continue

              image_bytes = decode_data_url(data_url)
              image_stream = io.BytesIO(image_bytes)

              p = doc.add_paragraph()
              p.alignment = WD_ALIGN_PARAGRAPH.CENTER
              run = p.add_run()
              run.add_picture(image_stream, width=Inches(5.8))

              name = str(item.get("name", "")).strip()
              if name:
                  cap = doc.add_paragraph()
                  cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                  cap.add_run(name)

      doc_stream = io.BytesIO()
      doc.save(doc_stream)
      doc_stream.seek(0)

      headers = {
          "Content-Disposition": 'attachment; filename="chat_export.docx"'
      }

      return StreamingResponse(
          doc_stream,
          media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
          headers=headers
      )
    except Exception as e:
      return JSONResponse(
          status_code=500,
          content={"error": str(e)}
      )